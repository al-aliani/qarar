import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ReportGenerator:
    """Generates professional Word documents for Feasibility Studies."""
    def __init__(self, project_name: str):
        self.project_name = project_name
        self.doc = Document()
        
    def add_title(self):
        title = self.doc.add_heading(f'دراسة جدوى: {self.project_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    def add_financial_summary(self, roi: float, breakeven_months: int):
        self.doc.add_heading('الملخص المالي', level=1)
        p = self.doc.add_paragraph()
        p.add_run(f'العائد المتوقع على الاستثمار (ROI): ').bold = True
        p.add_run(f'{roi}%\n')
        p.add_run(f'نقطة التعادل المتوقعة: ').bold = True
        p.add_run(f'{breakeven_months} أشهر')
        
    def save(self, filepath: str):
        self.doc.save(filepath)
        return filepath

if __name__ == "__main__":
    report = ReportGenerator("مشروع مقهى النرجس")
    report.add_title()
    report.add_financial_summary(roi=35.5, breakeven_months=14)
    output_path = os.path.join(os.getcwd(), "report_demo.docx")
    report.save(output_path)
    print(f"Report generated successfully at: {output_path}")
